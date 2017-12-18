import os
from docx import Document

directory = "/Users/tomshannon/Documents/GitHub/WSGC-Annual-Report-Mail-Merge-Tool/program_templates/"

target_document = Document()

for subdirectory in os.listdir(directory):
    
    if os.path.isdir(directory + subdirectory):
    
        for file in os.listdir(directory + subdirectory):

            target_document.add_page_break()

            for paragraph in Document(directory + subdirectory + "/" + file).paragraphs:

                text = paragraph.text

                target_document.add_paragraph(text)

            target_document.save("Final Report.docx")
